import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except Exception as exc:  # pragma: no cover
    raise SystemExit(
        "Missing dependency 'python-docx'. Install with: pip install python-docx"
    ) from exc


def add_title(document: Document, text: str) -> None:
    title = document.add_paragraph()
    run = title.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph(text)
    p_format = p.paragraph_format
    p_format.space_after = Pt(6)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_entities_table(document: Document, rows: list[dict]) -> None:
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Thực thể"
    hdr_cells[1].text = "Trường cốt lõi"
    hdr_cells[2].text = "Ghi chú"
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row.get("name", "")
        cells[1].text = ", ".join(row.get("fields", []))
        cells[2].text = row.get("note", "")


def build_document() -> Document:
    document = Document()

    add_title(document, "DNTU Focus - Mô tả bài toán và phân tích")
    add_paragraph(
        document,
        f"Tạo lúc: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )

    # 1) Tổng quan
    add_heading(document, "1. Tổng quan", level=1)
    add_paragraph(
        document,
        "DNTU Focus là ứng dụng quản lý công việc cá nhân kết hợp kỹ thuật Pomodoro, giúp người dùng tập trung, phân rã công việc và theo dõi tiến độ."
    )

    # 2) Mục tiêu & Phạm vi
    add_heading(document, "2. Mục tiêu", level=1)
    add_bullets(
        document,
        [
            "Quản lý danh sách công việc theo dự án",
            "Phân rã công việc thành các công việc con",
            "Theo dõi phiên làm việc Pomodoro",
            "Giao diện đơn giản, dễ thao tác trên di động",
        ],
    )

    add_heading(document, "3. Phạm vi", level=1)
    add_bullets(
        document,
        [
            "Quản lý người dùng (xác thực ngoài phạm vi tài liệu này)",
            "Quản lý dự án và công việc",
            "Quản lý công việc con (subtasks)",
            "Ghi nhận phiên Pomodoro tối thiểu (start/end)",
        ],
    )

    # 3) Yêu cầu chức năng
    add_heading(document, "4. Yêu cầu chức năng", level=1)
    add_numbered(
        document,
        [
            "Tạo/đọc/cập nhật/xoá (CRUD) dự án",
            "CRUD công việc: tiêu đề, gán dự án, thuộc người dùng",
            "CRUD công việc con: tiêu đề, trạng thái hoàn thành",
            "Bắt đầu/kết thúc phiên Pomodoro và lưu thời gian",
        ],
    )

    # 4) Yêu cầu phi chức năng
    add_heading(document, "5. Yêu cầu phi chức năng", level=1)
    add_bullets(
        document,
        [
            "Đơn giản, hiệu năng tốt cho danh sách vừa phải",
            "Khả năng mở rộng dần (tags, thông báo, đồng bộ cloud)",
            "Dữ liệu tách theo người dùng",
        ],
    )

    # 5) Mô hình dữ liệu (tối giản theo ERD hiện tại)
    add_heading(document, "6. Mô hình dữ liệu (tối giản)", level=1)
    add_paragraph(document, "Các thực thể cốt lõi và trường tối thiểu:")
    add_entities_table(
        document,
        rows=[
            {
                "name": "USER",
                "fields": ["id", "email"],
                "note": "Đại diện người dùng hệ thống",
            },
            {
                "name": "PROJECT",
                "fields": ["id", "name", "userId"],
                "note": "Nhóm công việc",
            },
            {
                "name": "TASK",
                "fields": ["id", "title", "projectId", "userId"],
                "note": "Thực thể công việc tối giản",
            },
            {
                "name": "SUBTASK",
                "fields": ["id", "title", "isCompleted", "taskId"],
                "note": "Công việc con",
            },
            {
                "name": "POMODORO_SESSION",
                "fields": ["id", "taskId", "startTime", "endTime"],
                "note": "Ghi nhận phiên làm việc",
            },
        ],
    )

    # 6) Luồng nghiệp vụ chính
    add_heading(document, "7. Luồng nghiệp vụ chính", level=1)
    add_numbered(
        document,
        [
            "Người dùng tạo Project",
            "Người dùng tạo Task thuộc Project",
            "(Tuỳ chọn) Người dùng tạo Subtask cho Task",
            "Người dùng bắt đầu Pomodoro, kết thúc và lưu phiên",
        ],
    )

    # 7) Khả năng mở rộng
    add_heading(document, "8. Khả năng mở rộng (tương lai)", level=1)
    add_bullets(
        document,
        [
            "Thêm TAG và quan hệ N-N với TASK",
            "Thông báo (local/cloud)",
            "Đồng bộ dữ liệu với Firestore",
            "Báo cáo hiệu suất (thống kê pomodoro, hoàn thành)",
        ],
    )

    # 9) Đối tượng sử dụng & Persona
    add_heading(document, "9. Đối tượng sử dụng & Persona", level=1)
    add_bullets(
        document,
        [
            "Sinh viên: cần quản lý môn học, deadline, ôn tập theo phiên Pomodoro",
            "Người đi làm: theo dõi task theo dự án, tập trung khi làm việc sâu",
            "Người tự học: chia nhỏ mục tiêu, duy trì thói quen học tập",
        ],
    )

    # 10) User Stories
    add_heading(document, "10. User Stories (ví dụ)", level=1)
    add_numbered(
        document,
        [
            "Là người dùng, tôi muốn tạo Task với tiêu đề để ghi lại việc cần làm",
            "Là người dùng, tôi muốn nhóm Task theo Project để dễ quản lý",
            "Là người dùng, tôi muốn tạo Subtask để chia nhỏ công việc phức tạp",
            "Là người dùng, tôi muốn bắt đầu/kết thúc Pomodoro để tập trung làm việc",
        ],
    )

    # 11) Use Cases tóm tắt
    add_heading(document, "11. Use Cases tóm tắt", level=1)
    add_bullets(
        document,
        [
            "UC-01 Tạo Project: nhập tên, lưu thuộc user",
            "UC-02 Tạo Task: nhập tiêu đề, gán vào Project",
            "UC-03 Quản lý Subtask: thêm/xoá/đánh dấu hoàn thành",
            "UC-04 Ghi nhận Pomodoro: start/end một phiên cho Task",
        ],
    )

    # 12) Ràng buộc & Giả định
    add_heading(document, "12. Ràng buộc & Giả định", level=1)
    add_bullets(
        document,
        [
            "Thiết bị mục tiêu: Android/iOS (ưu tiên di động)",
            "Lưu trữ local trước, đồng bộ cloud triển khai sau",
            "Đơn giản hoá mô hình dữ liệu để tăng tốc phát triển ban đầu",
        ],
    )

    # 13) Rủi ro & Giảm thiểu
    add_heading(document, "13. Rủi ro & Biện pháp giảm thiểu", level=1)
    add_bullets(
        document,
        [
            "Rủi ro phình tính năng: giữ phạm vi tối thiểu, lộ trình mở rộng",
            "Hiệu năng danh sách: phân trang/lazy load khi cần",
            "Mất dữ liệu local: sao lưu định kỳ, đồng bộ cloud sau này",
        ],
    )

    # 14) Kiến trúc mức cao
    add_heading(document, "14. Kiến trúc mức cao", level=1)
    add_paragraph(
        document,
        "Ứng dụng Flutter, tầng dữ liệu tối giản (USER/PROJECT/TASK/SUBTASK/POMODORO_SESSION). Có thể thêm repository pattern và đồng bộ Firestore ở giai đoạn 2.",
    )

    # 15) Quy trình Pomodoro (dòng chảy)
    add_heading(document, "15. Dòng chảy Pomodoro (cơ bản)", level=1)
    add_numbered(
        document,
        [
            "Chọn Task muốn tập trung",
            "Bắt đầu đếm Pomodoro (startTime)",
            "Làm việc tập trung đến khi kết thúc (endTime)",
            "Ghi nhận phiên Pomodoro vào lịch sử",
        ],
    )

    # 16) Backlog (ưu tiên)
    add_heading(document, "16. Backlog (ưu tiên)", level=1)
    add_numbered(
        document,
        [
            "MVP: CRUD Project/Task/Subtask, Pomodoro phiên (start/end)",
            "Bổ sung: Tag, lọc/sort, checklist nâng cao",
            "Thông báo: nhắc phiên/break (local)",
            "Đồng bộ Cloud Firestore, đăng nhập",
        ],
    )

    # 17) Tiêu chí chấp nhận (ví dụ)
    add_heading(document, "17. Tiêu chí chấp nhận (ví dụ)", level=1)
    add_bullets(
        document,
        [
            "Tạo Task với tiêu đề bắt buộc; thiếu tiêu đề -> báo lỗi",
            "Một Task có thể thêm nhiều Subtask và toggle hoàn thành",
            "Bắt đầu/kết thúc Pomodoro tạo được một bản ghi phiên",
        ],
    )

    # 18) Lộ trình
    add_heading(document, "18. Lộ trình", level=1)
    add_numbered(
        document,
        [
            "Giai đoạn 1: MVP tối giản (3-4 tuần)",
            "Giai đoạn 2: Đồng bộ cloud, thông báo, báo cáo (4-6 tuần)",
            "Giai đoạn 3: AI gợi ý, phân tích nâng cao (tuỳ chọn)",
        ],
    )

    # 19) Pomodoro là gì?
    add_heading(document, "19. Pomodoro là gì?", level=1)
    add_paragraph(
        document,
        "Pomodoro là kỹ thuật quản lý thời gian do Francesco Cirillo đề xuất cuối thập niên 1980. Tư tưởng chính: chia thời gian làm việc thành các phiên ngắn (thường 25 phút) gọi là một 'Pomodoro', đan xen với nghỉ ngắn (khoảng 5 phút) để duy trì sự tập trung cao và giảm mệt mỏi tinh thần.",
    )
    add_heading(document, "Mục tiêu của Pomodoro", level=2)
    add_bullets(
        document,
        [
            "Tăng thời gian tập trung sâu (deep work)",
            "Giảm phân tâm và trì hoãn",
            "Đo lường nỗ lực qua số phiên đã hoàn thành",
            "Tạo nhịp độ làm việc bền vững",
        ],
    )
    add_heading(document, "Quy tắc cơ bản", level=2)
    add_numbered(
        document,
        [
            "Chọn một công việc cụ thể",
            "Đặt hẹn giờ 25 phút và chỉ tập trung vào công việc đó",
            "Khi chuông kêu, nghỉ 5 phút",
            "Sau 4 phiên, nghỉ dài 15-30 phút",
        ],
    )
    add_heading(document, "Biến thể thường dùng", level=2)
    add_bullets(
        document,
        [
            "20/5: phù hợp với task ngắn",
            "30/5: phù hợp với task cần đà suy nghĩ",
            "50/10 (Ultradian): cho các công việc sáng tạo chuyên sâu",
        ],
    )
    add_heading(document, "Khi nào không nên dùng", level=2)
    add_bullets(
        document,
        [
            "Các cuộc họp hoặc lịch cố định khó chia nhỏ",
            "Công việc cần tương tác liên tục (support real-time)",
            "Tác vụ đòi hỏi thời gian liên tục dài (ví dụ build dài) — cân nhắc phiên dài hơn",
        ],
    )
    add_heading(document, "Cơ sở khoa học (tóm tắt)", level=2)
    add_paragraph(
        document,
        "Việc luân phiên giữa tập trung và nghỉ cho phép bộ não phục hồi sự chú ý (attentional resources), giảm hiện tượng suy giảm hiệu năng theo thời gian (vigilance decrement).",
    )

    # 20) Phân tích chi tiết yêu cầu chức năng
    add_heading(document, "20. Phân tích chi tiết yêu cầu chức năng", level=1)
    add_bullets(
        document,
        [
            "Project: tạo/đổi tên/xoá, xem danh sách project",
            "Task: tạo với tiêu đề bắt buộc, gán project, liệt kê theo project",
            "Subtask: thêm/xoá/đánh dấu hoàn thành, hiển thị theo task",
            "Pomodoro: start/end, thống kê số phiên theo task",
        ],
    )
    add_heading(document, "Luật nghiệp vụ bổ sung", level=2)
    add_bullets(
        document,
        [
            "Tiêu đề Task không rỗng",
            "Không cho xoá Project nếu vẫn còn Task (tuỳ lựa chọn)",
            "Một phiên Pomodoro phải có startTime < endTime",
        ],
    )

    # 21) Yêu cầu phi chức năng (mở rộng)
    add_heading(document, "21. Yêu cầu phi chức năng (mở rộng)", level=1)
    add_bullets(
        document,
        [
            "Khả dụng: thao tác mượt trên thiết bị tầm trung",
            "Khả mở rộng: dễ bổ sung tags/thông báo/đồng bộ cloud",
            "Khả bảo trì: cấu trúc mã rõ ràng, tách lớp UI/logic/data",
            "Khả quan sát: log sự kiện pomodoro và hành vi chính",
        ],
    )

    # 22) Kiến trúc & Dữ liệu (mở rộng)
    add_heading(document, "22. Kiến trúc & Dữ liệu (mở rộng)", level=1)
    add_paragraph(
        document,
        "Gợi ý kiến trúc: Flutter + Repository pattern; dữ liệu tối giản như ERD hiện tại. Có thể bổ sung cache local (Hive/SQLite) và nguồn đồng bộ (Firestore) ở bước sau.",
    )
    add_bullets(
        document,
        [
            "Tầng trình bày: Widget/State management (Provider/Bloc/Riverpod)",
            "Tầng domain: use-case (startPomodoro, completeSubtask...)",
            "Tầng data: repository -> datasource (local/cloud)",
        ],
    )

    # 23) Thuật ngữ & Chú giải
    add_heading(document, "23. Thuật ngữ & Chú giải", level=1)
    add_bullets(
        document,
        [
            "Pomodoro: một phiên tập trung có thời lượng cố định",
            "Subtask: công việc con của một Task",
            "Deep work: trạng thái tập trung cao độ trong thời gian đủ dài",
        ],
    )

    return document


def main() -> None:
    root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    out_dir = os.path.join(root, "docs")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "PhanTich_BaiToan_DNTU_Focus.docx")

    doc = build_document()
    doc.save(out_path)
    print(f"Generated: {out_path}")


if __name__ == "__main__":
    main()


